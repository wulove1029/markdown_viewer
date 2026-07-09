"""Tests for Markdown -> Word (.docx) export."""

import pytest

pytest.importorskip("docx")

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches

from app.docx_export import BODY_FONT, CODE_FONT, _content_width_emu, export_markdown_to_docx
from app.pptx_export import REMOTE_IMAGE_MAX_TOTAL_BYTES


def _text(doc):
    return "\n".join(p.text for p in doc.paragraphs)


def _east_asia_font(run):
    r_pr = run._element.rPr
    if r_pr is None or r_pr.rFonts is None:
        return None
    return r_pr.rFonts.get(qn("w:eastAsia"))


@pytest.fixture
def png_fixture(tmp_path):
    import pymupdf

    pix = pymupdf.Pixmap(pymupdf.csRGB, pymupdf.IRect(0, 0, 40, 24))
    pix.clear_with(210)
    path = tmp_path / "image.png"
    pix.save(str(path))
    return path


def test_export_roundtrip_heading_table_code_chinese_and_image(tmp_path, png_fixture):
    md = (
        "# 標題\n\n"
        "中文段落 with **bold** text.\n\n"
        "```py\nprint('hi')\n```\n\n"
        "| A | B |\n| - | - |\n| 1 | 2 |\n\n"
        f"![sample]({png_fixture.name})\n"
    )
    out = tmp_path / "doc.docx"

    count = export_markdown_to_docx(md, out, base_dir=tmp_path)

    assert count >= 5
    doc = Document(str(out))
    assert doc.paragraphs[0].style.name == "Heading 1"
    assert doc.paragraphs[0].text == "標題"
    assert any(run.text.startswith("中文段落") and _east_asia_font(run) == BODY_FONT
               for p in doc.paragraphs for run in p.runs)

    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert table.cell(0, 0).text == "A"
    assert table.cell(1, 1).text == "2"

    code_runs = [run for p in doc.paragraphs for run in p.runs if "print('hi')" in run.text]
    assert code_runs
    assert _east_asia_font(code_runs[0]) == CODE_FONT
    assert code_runs[0].font.name == CODE_FONT

    assert len(doc.inline_shapes) == 1


def test_table_cell_inline_formatting_is_preserved(tmp_path):
    md = (
        "| H1 | H2 | H3 |\n| - | - | - |\n"
        "| **bold** | *italic* | `code` |\n"
    )
    out = tmp_path / "table-inline.docx"

    export_markdown_to_docx(md, out)

    table = Document(str(out)).tables[0]
    bold_run = table.cell(1, 0).paragraphs[0].runs[0]
    italic_run = table.cell(1, 1).paragraphs[0].runs[0]
    code_run = table.cell(1, 2).paragraphs[0].runs[0]
    assert bold_run.bold is True
    assert italic_run.italic is True
    assert _east_asia_font(code_run) == CODE_FONT


def test_image_width_uses_pixels_and_dpi_without_upscaling(tmp_path):
    import pymupdf

    pix = pymupdf.Pixmap(pymupdf.csRGB, pymupdf.IRect(0, 0, 96, 48))
    pix.clear_with(180)
    pix.set_dpi(96, 96)
    small = tmp_path / "small.png"
    pix.save(str(small))
    out = tmp_path / "small.docx"

    export_markdown_to_docx(f"![small]({small.name})\n", out, base_dir=tmp_path)

    shape = Document(str(out)).inline_shapes[0]
    assert abs(shape.width - Inches(1)) < 3000


def test_large_image_clamps_to_docx_content_width(tmp_path):
    import pymupdf

    pix = pymupdf.Pixmap(pymupdf.csRGB, pymupdf.IRect(0, 0, 1800, 200))
    pix.clear_with(180)
    pix.set_dpi(96, 96)
    large = tmp_path / "large.png"
    pix.save(str(large))
    out = tmp_path / "large.docx"

    export_markdown_to_docx(f"![large]({large.name})\n", out, base_dir=tmp_path)

    doc = Document(str(out))
    assert doc.inline_shapes[0].width == _content_width_emu(doc)


def test_image_provider_embeds_mermaid_and_omits_source(tmp_path, png_fixture):
    md = "## Diagram\n\n```mermaid\ngraph TD; A-->B;\n```\n"
    out = tmp_path / "mermaid.docx"
    calls = []

    def provider(kind, source):
        calls.append((kind, source))
        return str(png_fixture)

    export_markdown_to_docx(md, out, image_provider=provider)
    doc = Document(str(out))

    assert calls == [("mermaid", "graph TD; A-->B;")]
    assert len(doc.inline_shapes) == 1
    assert "graph TD" not in _text(doc)


def test_provider_none_falls_back_to_source_text(tmp_path):
    md = "## Diagram\n\n```mermaid\ngraph LR; X-->Y;\n```\n"
    out = tmp_path / "fallback.docx"

    export_markdown_to_docx(md, out, image_provider=lambda kind, source: None)
    doc = Document(str(out))

    assert len(doc.inline_shapes) == 0
    assert "Mermaid diagram source" in _text(doc)
    assert "graph LR; X-->Y;" in _text(doc)


def test_missing_image_falls_back_to_placeholder(tmp_path):
    md = "![missing alt](missing.png)\n"
    out = tmp_path / "missing.docx"

    export_markdown_to_docx(md, out, base_dir=tmp_path)
    doc = Document(str(out))

    assert "[Image: missing alt]" in _text(doc)


def test_oversized_remote_image_falls_back_to_placeholder(tmp_path, monkeypatch):
    class FakeResponse:
        headers = {"Content-Length": str(REMOTE_IMAGE_MAX_TOTAL_BYTES + 1)}

        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, tb):
            return False

        def read(self, _size):
            raise AssertionError("oversized images should not be read")

    monkeypatch.setattr(
        "app.pptx_export.urllib.request.urlopen",
        lambda _url, timeout: FakeResponse(),
    )
    out = tmp_path / "remote.docx"

    export_markdown_to_docx("![remote](https://example.test/huge.png)\n", out)
    doc = Document(str(out))

    assert len(doc.inline_shapes) == 0
    assert "[Image: remote]" in _text(doc)
