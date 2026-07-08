"""Export a Markdown document to an editable Word (.docx) document.

The Markdown block model is shared with the PowerPoint exporter so headings,
tables, lists, code blocks, images, Mermaid, and math fragments are parsed in
one place. Rendering here is intentionally Qt-free; the GUI may pass an
``image_provider(kind, source) -> png_path | None`` callback for Mermaid/math.
"""

from __future__ import annotations

import io
import urllib.request
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from .pptx_export import (
    Code,
    Heading,
    Image,
    ListBlock,
    Para,
    Quote,
    Run,
    Table,
    parse_elements,
)

BODY_FONT = "Microsoft JhengHei"
CODE_FONT = "Consolas"
BODY_PT = 11
CODE_PT = 9
IMAGE_WIDTH = Inches(6.2)


def _rfonts(run_or_style, font_name: str) -> None:
    """Set ASCII/HAnsi and East Asian font names on a run or style."""
    if hasattr(run_or_style, "font"):
        run_or_style.font.name = font_name
    element = getattr(run_or_style, "_element", None)
    if element is None:
        element = getattr(run_or_style, "_r", None)
    if element is None:
        return
    if hasattr(element, "get_or_add_rPr"):
        r_pr = element.get_or_add_rPr()
    else:
        r_pr = getattr(element, "rPr", None)
        if r_pr is None:
            r_pr = OxmlElement("w:rPr")
            element.append(r_pr)
    r_fonts = getattr(r_pr, "rFonts", None)
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        r_fonts.set(qn(key), font_name)


def _set_doc_defaults(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.size = Pt(BODY_PT)
    _rfonts(normal, BODY_FONT)
    for style_name in ("Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        if style_name in doc.styles:
            _rfonts(doc.styles[style_name], BODY_FONT)


def _style_exists(doc: Document, name: str) -> bool:
    try:
        doc.styles[name]
    except KeyError:
        return False
    return True


def _list_style(doc: Document, ordered: bool, level: int) -> str | None:
    base = "List Number" if ordered else "List Bullet"
    candidates = [base] if level <= 0 else [f"{base} {min(level + 1, 3)}", base]
    for name in candidates:
        if _style_exists(doc, name):
            return name
    return None


def _text_of(runs: list[Run]) -> str:
    return "".join(r.text for r in runs)


def _emit_runs(paragraph, runs: list[Run], *, size_pt: int = BODY_PT) -> None:
    if not runs:
        runs = [Run("")]
    for source in runs:
        run = paragraph.add_run(source.text)
        run.bold = source.bold
        run.italic = source.italic
        run.font.size = Pt(size_pt)
        if source.code:
            _rfonts(run, CODE_FONT)
        else:
            _rfonts(run, BODY_FONT)


def _add_heading(doc: Document, block: Heading) -> None:
    level = max(1, min(9, block.level))
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    _emit_runs(paragraph, block.runs, size_pt=max(12, 22 - level * 2))


def _add_paragraph(doc: Document, runs: list[Run], style: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    _emit_runs(paragraph, runs)


def _shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def _add_code(doc: Document, block: Code, base_dir, image_provider=None) -> None:
    if block.lang in ("mermaid", "math") and image_provider is not None:
        rendered = image_provider(block.lang, block.text)
        if rendered and _add_image(
            doc, Image(str(rendered), block.lang), None, placeholder=False
        ):
            return

    if block.lang in ("mermaid", "math"):
        label = "Mermaid diagram source" if block.lang == "mermaid" else "LaTeX math source"
        label_para = doc.add_paragraph()
        label_run = label_para.add_run(label)
        label_run.italic = True
        _rfonts(label_run, BODY_FONT)

    paragraph = doc.add_paragraph()
    _shade_paragraph(paragraph, "F2F2F2")
    run = paragraph.add_run(block.text or "")
    run.font.size = Pt(CODE_PT)
    _rfonts(run, CODE_FONT)


def _add_table(doc: Document, block: Table) -> None:
    n_cols = max([len(block.header)] + [len(row) for row in block.rows] or [1]) or 1
    data_rows = ([block.header] if block.header else []) + block.rows
    if not data_rows:
        data_rows = [[""]]
    table = doc.add_table(rows=len(data_rows), cols=n_cols)
    table.style = "Table Grid"
    for row_idx, row in enumerate(data_rows):
        for col_idx in range(n_cols):
            cell = table.cell(row_idx, col_idx)
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(row[col_idx] if col_idx < len(row) else "")
            if block.header and row_idx == 0:
                run.bold = True
            _rfonts(run, BODY_FONT)
    doc.add_paragraph()


def _load_image_stream(src: str, base_dir):
    if src.startswith(("http://", "https://")):
        try:
            with urllib.request.urlopen(src, timeout=8) as resp:  # nosec - user content
                return io.BytesIO(resp.read())
        except Exception:
            return None
    path = Path(src)
    if not path.is_absolute() and base_dir is not None:
        path = Path(base_dir) / src
    try:
        return io.BytesIO(path.read_bytes())
    except Exception:
        return None


def _add_image(doc: Document, block: Image, base_dir, *, placeholder: bool = True) -> bool:
    stream = _load_image_stream(block.src, base_dir)
    if stream is None:
        if placeholder:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(f"[Image: {block.alt or block.src}]")
            run.italic = True
            _rfonts(run, BODY_FONT)
        return False
    try:
        doc.add_picture(stream, width=IMAGE_WIDTH)
    except Exception:
        if placeholder:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(f"[Image: {block.alt or block.src}]")
            run.italic = True
            _rfonts(run, BODY_FONT)
        return False
    return True


def export_markdown_to_docx(md_text: str, out_path, base_dir=None, image_provider=None) -> int:
    """Convert *md_text* to a .docx document at *out_path*.

    Returns the number of parsed Markdown block elements rendered. Mermaid and
    math blocks are embedded as PNG images when *image_provider* succeeds;
    otherwise their source text is kept in a code block.
    """
    elements = parse_elements(md_text)
    doc = Document()
    _set_doc_defaults(doc)

    for block in elements:
        if isinstance(block, Heading):
            _add_heading(doc, block)
        elif isinstance(block, Para):
            _add_paragraph(doc, block.runs)
        elif isinstance(block, Quote):
            style = "Quote" if _style_exists(doc, "Quote") else None
            _add_paragraph(doc, block.runs, style=style)
        elif isinstance(block, ListBlock):
            for item in block.items:
                style = _list_style(doc, item.ordered, item.level)
                paragraph = doc.add_paragraph(style=style)
                if style is None:
                    prefix = f"{item.number}. " if item.ordered else "- "
                    paragraph.add_run(("    " * item.level) + prefix)
                _emit_runs(paragraph, item.runs)
        elif isinstance(block, Code):
            _add_code(doc, block, base_dir, image_provider)
        elif isinstance(block, Table):
            _add_table(doc, block)
        elif isinstance(block, Image):
            _add_image(doc, block, base_dir)

    if not elements:
        doc.add_paragraph()

    doc.save(str(out_path))
    return len(elements)
